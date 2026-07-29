import csv
from collections.abc import Iterable
from io import BytesIO, StringIO
from pathlib import Path
from typing import Protocol
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
from ai_api.rag.exceptions import TextExtractionError
from ai_api.rag.schemas import TextExtractionResponse


class FileExtractor(Protocol):
    name: str
    supported_extensions: frozenset[str]

    def extract(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TextExtractionResponse:
        """Extract normalized content from a supported file."""
        ...


def normalize_extracted_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def normalize_table_cell(value: object) -> str:
    if value is None:
        return ""

    return " ".join(
        normalize_extracted_text(str(value)).split()
    )


def get_file_extension(filename: str) -> str:
    cleaned_filename = filename.strip() or "uploaded-file"

    return Path(cleaned_filename).suffix.lower()


class Utf8TextFileExtractor:
    name = "utf-8-text"

    supported_extensions = frozenset(
        {
            ".txt",
            ".md",
            ".markdown",
        }
    )

    def extract(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TextExtractionResponse:
        cleaned_filename = filename.strip() or "uploaded-file"
        extension = get_file_extension(cleaned_filename)

        if extension not in self.supported_extensions:
            raise TextExtractionError(
                f"Unsupported file type: {extension or 'unknown'}"
            )

        try:
            extracted_text = file_content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise TextExtractionError(
                "File content could not be decoded as UTF-8 text."
            ) from exc

        normalized_text = normalize_extracted_text(extracted_text)

        if not normalized_text:
            raise TextExtractionError("Extracted text is empty.")

        return TextExtractionResponse(
            source=cleaned_filename,
            filename=cleaned_filename,
            content_type=content_type,
            character_count=len(normalized_text),
            text=normalized_text,
            metadata={
                "extension": extension,
                "extraction_method": self.name,
                "extractor": type(self).__name__,
            },
        )


class PDFFileExtractor:
    name = "pdf"

    supported_extensions = frozenset(
        {
            ".pdf",
        }
    )

    def extract(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TextExtractionResponse:
        cleaned_filename = filename.strip() or "uploaded-file"
        extension = get_file_extension(cleaned_filename)

        if extension not in self.supported_extensions:
            raise TextExtractionError(
                f"Unsupported file type: {extension or 'unknown'}"
            )

        try:
            reader = PdfReader(BytesIO(file_content))

            if reader.is_encrypted:
                raise TextExtractionError(
                    "Encrypted PDF files are not supported."
                )

            page_texts = [
                normalize_extracted_text(page.extract_text() or "")
                for page in reader.pages
            ]

            page_count = len(reader.pages)
        except TextExtractionError:
            raise
        except Exception as exc:
            raise TextExtractionError(
                "PDF content could not be extracted."
            ) from exc

        normalized_text = normalize_extracted_text(
            "\n\n".join(
                page_text
                for page_text in page_texts
                if page_text
            )
        )

        if not normalized_text:
            raise TextExtractionError("Extracted text is empty.")

        return TextExtractionResponse(
            source=cleaned_filename,
            filename=cleaned_filename,
            content_type=content_type,
            character_count=len(normalized_text),
            text=normalized_text,
            metadata={
                "extension": extension,
                "extraction_method": self.name,
                "extractor": type(self).__name__,
                "page_count": page_count,
            },
        )


class DOCXFileExtractor:
    name = "docx"

    supported_extensions = frozenset(
        {
            ".docx",
        }
    )

    def extract(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TextExtractionResponse:
        cleaned_filename = filename.strip() or "uploaded-file"
        extension = get_file_extension(cleaned_filename)

        if extension not in self.supported_extensions:
            raise TextExtractionError(
                f"Unsupported file type: {extension or 'unknown'}"
            )

        try:
            document = Document(BytesIO(file_content))
        except Exception as exc:
            raise TextExtractionError(
                "DOCX content could not be extracted."
            ) from exc

        paragraph_texts = [
            normalize_extracted_text(paragraph.text)
            for paragraph in document.paragraphs
            if normalize_extracted_text(paragraph.text)
        ]

        table_texts: list[str] = []

        for table in document.tables:
            for row in table.rows:
                row_values = [
                    normalize_table_cell(cell.text)
                    for cell in row.cells
                    if normalize_table_cell(cell.text)
                ]

                if row_values:
                    table_texts.append(" | ".join(row_values))

        normalized_text = normalize_extracted_text(
            "\n\n".join(
                [
                    *paragraph_texts,
                    *table_texts,
                ]
            )
        )

        if not normalized_text:
            raise TextExtractionError("Extracted text is empty.")

        return TextExtractionResponse(
            source=cleaned_filename,
            filename=cleaned_filename,
            content_type=content_type,
            character_count=len(normalized_text),
            text=normalized_text,
            metadata={
                "extension": extension,
                "extraction_method": self.name,
                "extractor": type(self).__name__,
                "paragraph_count": len(paragraph_texts),
                "table_count": len(document.tables),
            },
        )


class CSVFileExtractor:
    name = "csv"

    supported_extensions = frozenset(
        {
            ".csv",
        }
    )

    def extract(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TextExtractionResponse:
        cleaned_filename = filename.strip() or "uploaded-file"
        extension = get_file_extension(cleaned_filename)

        if extension not in self.supported_extensions:
            raise TextExtractionError(
                f"Unsupported file type: {extension or 'unknown'}"
            )

        try:
            decoded_content = file_content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise TextExtractionError(
                "CSV content could not be decoded as UTF-8 text."
            ) from exc

        if not normalize_extracted_text(decoded_content):
            raise TextExtractionError("Extracted text is empty.")

        try:
            dialect = csv.Sniffer().sniff(decoded_content[:2048])
        except csv.Error:
            dialect = csv.excel

        rows: list[list[str]] = []

        reader = csv.reader(
            StringIO(decoded_content),
            dialect=dialect,
        )

        for row in reader:
            normalized_row = [
                normalize_table_cell(value)
                for value in row
            ]

            while normalized_row and not normalized_row[-1]:
                normalized_row.pop()

            if any(normalized_row):
                rows.append(normalized_row)

        if not rows:
            raise TextExtractionError("Extracted text is empty.")

        normalized_text = normalize_extracted_text(
            "\n".join(
                " | ".join(row)
                for row in rows
            )
        )

        column_count = max(
            len(row)
            for row in rows
        )

        return TextExtractionResponse(
            source=cleaned_filename,
            filename=cleaned_filename,
            content_type=content_type,
            character_count=len(normalized_text),
            text=normalized_text,
            metadata={
                "extension": extension,
                "extraction_method": self.name,
                "extractor": type(self).__name__,
                "row_count": len(rows),
                "column_count": column_count,
                "delimiter": dialect.delimiter,
            },
        )


class ExcelFileExtractor:
    name = "excel"

    supported_extensions = frozenset(
        {
            ".xlsx",
        }
    )

    def extract(
        self,
        file_content: bytes,
        filename: str,
        content_type: str | None = None,
    ) -> TextExtractionResponse:
        cleaned_filename = filename.strip() or "uploaded-file"
        extension = get_file_extension(cleaned_filename)

        if extension not in self.supported_extensions:
            raise TextExtractionError(
                f"Unsupported file type: {extension or 'unknown'}"
            )

        try:
            workbook = load_workbook(
                BytesIO(file_content),
                read_only=True,
                data_only=True,
            )
        except Exception as exc:
            raise TextExtractionError(
                "Excel content could not be extracted."
            ) from exc

        try:
            sheet_names = workbook.sheetnames
            sheet_texts: list[str] = []
            row_count = 0
            column_count = 0

            for worksheet in workbook.worksheets:
                rows: list[list[str]] = []

                for row in worksheet.iter_rows(values_only=True):
                    normalized_row = [
                        normalize_table_cell(value)
                        for value in row
                    ]

                    while normalized_row and not normalized_row[-1]:
                        normalized_row.pop()

                    if any(normalized_row):
                        rows.append(normalized_row)

                if not rows:
                    continue

                row_count += len(rows)

                column_count = max(
                    column_count,
                    max(
                        len(row)
                        for row in rows
                    ),
                )

                sheet_texts.append(
                    normalize_extracted_text(
                        "\n".join(
                            [
                                f"# Sheet: {worksheet.title}",
                                *[
                                    " | ".join(row)
                                    for row in rows
                                ],
                            ]
                        )
                    )
                )

            normalized_text = normalize_extracted_text(
                "\n\n".join(sheet_texts)
            )

            if not normalized_text:
                raise TextExtractionError("Extracted text is empty.")

            return TextExtractionResponse(
                source=cleaned_filename,
                filename=cleaned_filename,
                content_type=content_type,
                character_count=len(normalized_text),
                text=normalized_text,
                metadata={
                    "extension": extension,
                    "extraction_method": self.name,
                    "extractor": type(self).__name__,
                    "sheet_count": len(sheet_names),
                    "sheet_names": sheet_names,
                    "row_count": row_count,
                    "column_count": column_count,
                },
            )
        finally:
            workbook.close()


class FileExtractorRegistry:
    def __init__(
        self,
        extractors: Iterable[FileExtractor] | None = None,
    ) -> None:
        selected_extractors = (
            [
                Utf8TextFileExtractor(),
                PDFFileExtractor(),
                DOCXFileExtractor(),
                CSVFileExtractor(),
                ExcelFileExtractor(),
            ]
            if extractors is None
            else list(extractors)
        )

        self._extractors_by_extension: dict[str, FileExtractor] = {}

        for extractor in selected_extractors:
            self.register(extractor)

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return tuple(sorted(self._extractors_by_extension))

    def register(self, extractor: FileExtractor) -> None:
        if not extractor.supported_extensions:
            raise ValueError(
                "File extractor must define at least one supported extension."
            )

        normalized_extensions = {
            self._normalize_extension(extension)
            for extension in extractor.supported_extensions
        }

        for extension in normalized_extensions:
            if extension in self._extractors_by_extension:
                raise ValueError(
                    f"Extractor already registered for extension: {extension}"
                )

        for extension in normalized_extensions:
            self._extractors_by_extension[extension] = extractor

    def get_for_filename(self, filename: str) -> FileExtractor:
        cleaned_filename = filename.strip() or "uploaded-file"
        extension = Path(cleaned_filename).suffix.lower()

        return self.get_for_extension(extension)

    def get_for_extension(self, extension: str) -> FileExtractor:
        normalized_extension = self._normalize_extension(
            extension,
            allow_empty=True,
        )

        extractor = self._extractors_by_extension.get(
            normalized_extension
        )

        if extractor is None:
            raise TextExtractionError(
                "Unsupported file type: "
                f"{normalized_extension or 'unknown'}"
            )

        return extractor

    def _normalize_extension(
        self,
        extension: str,
        allow_empty: bool = False,
    ) -> str:
        normalized_extension = extension.strip().lower()

        if not normalized_extension:
            if allow_empty:
                return ""

            raise ValueError("File extension cannot be blank.")

        if not normalized_extension.startswith("."):
            normalized_extension = f".{normalized_extension}"

        return normalized_extension
