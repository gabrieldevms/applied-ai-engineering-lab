from io import BytesIO
import pytest
from docx import Document
from reportlab.pdfgen import canvas
from ai_api.rag.exceptions import TextExtractionError
from ai_api.rag.file_extractors import DOCXFileExtractor, PDFFileExtractor
from ai_api.rag.text_extraction import TextExtractionService


def _build_pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()

    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, text)
    pdf.save()

    return buffer.getvalue()


def _build_docx_bytes() -> bytes:
    buffer = BytesIO()

    document = Document()
    document.add_heading("Requirement", level=1)
    document.add_paragraph(
        "Como cliente, quero gerar boleto para pagar minha dívida."
    )

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "status"
    table.cell(1, 1).text = "active"

    document.save(buffer)

    return buffer.getvalue()


def test_pdf_file_extractor_should_extract_text_from_pdf() -> None:
    extractor = PDFFileExtractor()

    response = extractor.extract(
        file_content=_build_pdf_bytes(
            "Como cliente, quero gerar boleto."
        ),
        filename="requirement.pdf",
        content_type="application/pdf",
    )

    assert response.filename == "requirement.pdf"
    assert "Como cliente, quero gerar boleto." in response.text
    assert response.content_type == "application/pdf"
    assert response.metadata["extension"] == ".pdf"
    assert response.metadata["extraction_method"] == "pdf"
    assert response.metadata["extractor"] == "PDFFileExtractor"
    assert response.metadata["page_count"] == 1


def test_docx_file_extractor_should_extract_paragraphs_and_tables() -> None:
    extractor = DOCXFileExtractor()

    response = extractor.extract(
        file_content=_build_docx_bytes(),
        filename="requirement.docx",
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    assert response.filename == "requirement.docx"
    assert "Requirement" in response.text
    assert "Como cliente, quero gerar boleto" in response.text
    assert "Field | Value" in response.text
    assert "status | active" in response.text
    assert response.metadata["extension"] == ".docx"
    assert response.metadata["extraction_method"] == "docx"
    assert response.metadata["extractor"] == "DOCXFileExtractor"
    assert response.metadata["paragraph_count"] == 2
    assert response.metadata["table_count"] == 1


def test_text_extraction_service_should_extract_pdf_with_default_registry() -> None:
    service = TextExtractionService()

    response = service.extract_from_bytes(
        file_content=_build_pdf_bytes(
            "Requisito extraído de PDF."
        ),
        filename="sample.pdf",
        content_type="application/pdf",
    )

    assert "Requisito extraído de PDF." in response.text
    assert response.metadata["extractor"] == "PDFFileExtractor"


def test_text_extraction_service_should_extract_docx_with_default_registry() -> None:
    service = TextExtractionService()

    response = service.extract_from_bytes(
        file_content=_build_docx_bytes(),
        filename="sample.docx",
    )

    assert "Como cliente, quero gerar boleto" in response.text
    assert response.metadata["extractor"] == "DOCXFileExtractor"


def test_pdf_file_extractor_should_reject_invalid_pdf_content() -> None:
    extractor = PDFFileExtractor()

    with pytest.raises(
        TextExtractionError,
        match="PDF content could not be extracted.",
    ):
        extractor.extract(
            file_content=b"not a real pdf",
            filename="broken.pdf",
            content_type="application/pdf",
        )


def test_docx_file_extractor_should_reject_invalid_docx_content() -> None:
    extractor = DOCXFileExtractor()

    with pytest.raises(
        TextExtractionError,
        match="DOCX content could not be extracted.",
    ):
        extractor.extract(
            file_content=b"not a real docx",
            filename="broken.docx",
        )
