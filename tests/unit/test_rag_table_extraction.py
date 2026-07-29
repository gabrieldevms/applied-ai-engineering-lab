from io import BytesIO
import pytest
from docx import Document
from openpyxl import Workbook
from ai_api.rag.exceptions import TextExtractionError
from ai_api.rag.table_extraction import (
    CSVStructuredTableExtractor,
    DOCXStructuredTableExtractor,
    ExcelStructuredTableExtractor,
    StructuredTableExtractorRegistry,
    TableExtractionService,
)


def _build_xlsx_bytes() -> bytes:
    buffer = BytesIO()

    workbook = Workbook()

    worksheet = workbook.active
    worksheet.title = "Requirements"
    worksheet.append(["Field", "Value"])
    worksheet.append(["status", "active"])

    rules_sheet = workbook.create_sheet("Rules")
    rules_sheet.append(["rule", "boleto obrigatório"])

    workbook.save(buffer)

    return buffer.getvalue()


def _build_docx_bytes_with_table() -> bytes:
    buffer = BytesIO()

    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "status"
    table.cell(1, 1).text = "active"

    document.save(buffer)

    return buffer.getvalue()


def test_structured_table_registry_should_register_default_extensions() -> None:
    registry = StructuredTableExtractorRegistry()

    assert registry.supported_extensions == (
        ".csv",
        ".docx",
        ".xlsx",
    )


def test_structured_table_registry_should_resolve_extractors_by_extension() -> None:
    registry = StructuredTableExtractorRegistry()

    assert isinstance(
        registry.get_for_filename("data.csv"),
        CSVStructuredTableExtractor,
    )
    assert isinstance(
        registry.get_for_filename("requirements.xlsx"),
        ExcelStructuredTableExtractor,
    )
    assert isinstance(
        registry.get_for_filename("document.docx"),
        DOCXStructuredTableExtractor,
    )


def test_table_extraction_service_should_extract_csv_table() -> None:
    service = TableExtractionService()

    response = service.extract_from_bytes(
        file_content=(
            "Field,Value\n"
            "status,active\n"
        ).encode("utf-8"),
        filename="data.csv",
        content_type="text/csv",
    )

    assert response.filename == "data.csv"
    assert response.content_type == "text/csv"
    assert response.table_count == 1

    table = response.tables[0]

    assert table.filename == "data.csv"
    assert table.rows == [
        ["Field", "Value"],
        ["status", "active"],
    ]
    assert table.row_count == 2
    assert table.column_count == 2
    assert table.metadata["delimiter"] == ","


def test_table_extraction_service_should_extract_excel_tables_by_sheet() -> None:
    service = TableExtractionService()

    response = service.extract_from_bytes(
        file_content=_build_xlsx_bytes(),
        filename="requirements.xlsx",
    )

    assert response.filename == "requirements.xlsx"
    assert response.table_count == 2
    assert response.metadata["sheet_count"] == 2
    assert response.metadata["sheet_names"] == [
        "Requirements",
        "Rules",
    ]

    first_table = response.tables[0]
    second_table = response.tables[1]

    assert first_table.sheet_name == "Requirements"
    assert first_table.rows == [
        ["Field", "Value"],
        ["status", "active"],
    ]

    assert second_table.sheet_name == "Rules"
    assert second_table.rows == [
        ["rule", "boleto obrigatório"],
    ]


def test_table_extraction_service_should_extract_docx_tables() -> None:
    service = TableExtractionService()

    response = service.extract_from_bytes(
        file_content=_build_docx_bytes_with_table(),
        filename="requirement.docx",
    )

    assert response.filename == "requirement.docx"
    assert response.table_count == 1

    table = response.tables[0]

    assert table.rows == [
        ["Field", "Value"],
        ["status", "active"],
    ]
    assert table.row_count == 2
    assert table.column_count == 2


def test_table_extraction_service_should_reject_unsupported_table_file_type() -> None:
    service = TableExtractionService()

    with pytest.raises(
        TextExtractionError,
        match="Unsupported table extraction file type: .pdf",
    ):
        service.extract_from_bytes(
            file_content=b"fake pdf content",
            filename="document.pdf",
            content_type="application/pdf",
        )


def test_excel_table_extractor_should_reject_invalid_xlsx_content() -> None:
    extractor = ExcelStructuredTableExtractor()

    with pytest.raises(
        TextExtractionError,
        match="Excel tables could not be extracted.",
    ):
        extractor.extract_tables(
            file_content=b"not a real xlsx file",
            filename="broken.xlsx",
        )
